import json
import pdfplumber
from io import BytesIO

from django.http import HttpResponse
from rest_framework import status
from rest_framework.decorators import api_view
from rest_framework.response import Response

from .models import (ResumeProfile, WorkExperience, Education,
                     Project, Skill, Certification, ResumeVersion, ATSResult)
from .serializers import (ResumeProfileSerializer, ResumeProfileCreateSerializer,
                           WorkExperienceSerializer, EducationSerializer,
                           ProjectSerializer, SkillSerializer,
                           CertificationSerializer, ResumeVersionSerializer,
                           ATSResultSerializer)
from .services.groq_service import (improve_bullet, categorize_skills,
                                     jd_match_analysis, parse_resume_text,
                                     generate_summary)


# ── RESUME PROFILE CRUD ────────────────────────────────────────────────────

@api_view(['GET', 'POST'])
def resume_list(request):
    if request.method == 'GET':
        resumes = ResumeProfile.objects.all().order_by('-created_at')
        return Response(ResumeProfileSerializer(resumes, many=True).data)

    serializer = ResumeProfileCreateSerializer(data=request.data)
    if serializer.is_valid():
        serializer.save()
        return Response(serializer.data, status=status.HTTP_201_CREATED)
    return Response(serializer.errors, status=status.HTTP_400_BAD_REQUEST)


@api_view(['GET', 'PUT', 'DELETE'])
def resume_detail(request, pk):
    try:
        resume = ResumeProfile.objects.get(pk=pk)
    except ResumeProfile.DoesNotExist:
        return Response({'error': 'Resume not found'}, status=404)

    if request.method == 'GET':
        return Response(ResumeProfileSerializer(resume).data)

    if request.method == 'PUT':
        serializer = ResumeProfileCreateSerializer(resume, data=request.data, partial=True)
        if serializer.is_valid():
            serializer.save()
            return Response(serializer.data)
        return Response(serializer.errors, status=400)

    resume.delete()
    return Response({'message': 'Deleted'}, status=204)


# ── WORK EXPERIENCE ────────────────────────────────────────────────────────

@api_view(['POST'])
def add_work_experience(request, resume_pk):
    try:
        resume = ResumeProfile.objects.get(pk=resume_pk)
    except ResumeProfile.DoesNotExist:
        return Response({'error': 'Resume not found'}, status=404)

    data = request.data.copy()
    data['resume'] = resume.id
    serializer = WorkExperienceSerializer(data=data)
    if serializer.is_valid():
        serializer.save(resume=resume)
        return Response(serializer.data, status=201)
    return Response(serializer.errors, status=400)


@api_view(['PUT', 'DELETE'])
def work_experience_detail(request, pk):
    try:
        exp = WorkExperience.objects.get(pk=pk)
    except WorkExperience.DoesNotExist:
        return Response({'error': 'Not found'}, status=404)

    if request.method == 'PUT':
        serializer = WorkExperienceSerializer(exp, data=request.data, partial=True)
        if serializer.is_valid():
            serializer.save()
            return Response(serializer.data)
        return Response(serializer.errors, status=400)

    exp.delete()
    return Response({'message': 'Deleted'}, status=204)


# ── EDUCATION ──────────────────────────────────────────────────────────────

@api_view(['POST'])
def add_education(request, resume_pk):
    try:
        resume = ResumeProfile.objects.get(pk=resume_pk)
    except ResumeProfile.DoesNotExist:
        return Response({'error': 'Resume not found'}, status=404)

    serializer = EducationSerializer(data=request.data)
    if serializer.is_valid():
        serializer.save(resume=resume)
        return Response(serializer.data, status=201)
    return Response(serializer.errors, status=400)


# ── PROJECTS ───────────────────────────────────────────────────────────────

@api_view(['POST'])
def add_project(request, resume_pk):
    try:
        resume = ResumeProfile.objects.get(pk=resume_pk)
    except ResumeProfile.DoesNotExist:
        return Response({'error': 'Resume not found'}, status=404)

    serializer = ProjectSerializer(data=request.data)
    if serializer.is_valid():
        serializer.save(resume=resume)
        return Response(serializer.data, status=201)
    return Response(serializer.errors, status=400)


# ── SKILLS ─────────────────────────────────────────────────────────────────

@api_view(['POST'])
def add_skill(request, resume_pk):
    try:
        resume = ResumeProfile.objects.get(pk=resume_pk)
    except ResumeProfile.DoesNotExist:
        return Response({'error': 'Resume not found'}, status=404)

    serializer = SkillSerializer(data=request.data)
    if serializer.is_valid():
        serializer.save(resume=resume)
        return Response(serializer.data, status=201)
    return Response(serializer.errors, status=400)


# ── AI ENDPOINTS ───────────────────────────────────────────────────────────

@api_view(['POST'])
def ai_improve_bullet(request):
    """
    Body: { "bullet": "worked on backend", "target_role": "Senior SWE at Google" }
    """
    bullet      = request.data.get('bullet', '')
    target_role = request.data.get('target_role', '')

    if not bullet:
        return Response({'error': 'bullet is required'}, status=400)

    result = improve_bullet(bullet, target_role)
    return Response(result)


@api_view(['POST'])
def ai_categorize_skills(request):
    """
    Body: { "skills": ["Python", "React", "Docker", "Leadership"] }
    """
    skills = request.data.get('skills', [])
    if not skills:
        return Response({'error': 'skills list is required'}, status=400)

    result = categorize_skills(skills)
    return Response(result)


@api_view(['POST'])
def ai_jd_match(request, resume_pk):
    """
    Body: { "job_description": "We need a Python engineer..." }
    Saves result to ATSResult table + returns score.
    """
    try:
        resume = ResumeProfile.objects.get(pk=resume_pk)
    except ResumeProfile.DoesNotExist:
        return Response({'error': 'Resume not found'}, status=404)

    jd = request.data.get('job_description', '')
    if not jd:
        return Response({'error': 'job_description is required'}, status=400)

    # Build resume text from all sections
    bullets = []
    for exp in resume.work_experiences.all():
        bullets.extend(exp.bullets)

    skill_names  = list(resume.skills.values_list('name', flat=True))
    resume_text  = ' '.join(filter(None, [
        resume.summary,
        ' '.join(bullets),
        ' '.join(skill_names)
    ]))

    result = jd_match_analysis(resume_text, jd)

    # Save to DB
    ats = ATSResult.objects.create(
        resume          = resume,
        job_description = jd,
        ats_score       = result.get('ats_score', 0),
        matched_keywords = result.get('matched_keywords', []),
        missing_keywords = result.get('missing_keywords', []),
        feedback        = result.get('feedback', '')
    )

    return Response(ATSResultSerializer(ats).data)


@api_view(['POST'])
def ai_generate_summary(request, resume_pk):
    """Auto-generates a professional summary for an existing resume."""
    try:
        resume = ResumeProfile.objects.get(pk=resume_pk)
    except ResumeProfile.DoesNotExist:
        return Response({'error': 'Resume not found'}, status=404)

    profile_data = ResumeProfileSerializer(resume).data
    summary = generate_summary(profile_data, resume.target_role)

    resume.summary = summary
    resume.save()
    return Response({'summary': summary})


# ── UPLOAD & PARSE ─────────────────────────────────────────────────────────

@api_view(['POST'])
def upload_and_parse(request):
    """
    Upload PDF or DOCX → AI parses → returns structured JSON.
    Does NOT save automatically — frontend confirms first.
    Form-data: file=<file>, target_role=<string>
    """
    file = request.FILES.get('file')
    if not file:
        return Response({'error': 'No file uploaded'}, status=400)

    target_role = request.data.get('target_role', '')

    # Extract raw text
    raw_text = ''
    if file.name.endswith('.pdf'):
        try:
            with pdfplumber.open(BytesIO(file.read())) as pdf:
                raw_text = '\n'.join(
                    page.extract_text() or '' for page in pdf.pages
                )
        except Exception as e:
            return Response({'error': f'PDF parse failed: {e}'}, status=400)

    elif file.name.endswith('.docx'):
        try:
            from docx import Document
            doc      = Document(BytesIO(file.read()))
            raw_text = '\n'.join(p.text for p in doc.paragraphs)
        except Exception as e:
            return Response({'error': f'DOCX parse failed: {e}'}, status=400)

    elif file.name.endswith('.txt'):
        raw_text = file.read().decode('utf-8')

    else:
        return Response({'error': 'Only PDF, DOCX, TXT supported'}, status=400)

    if not raw_text.strip():
        return Response({'error': 'Could not extract text from file'}, status=400)

    parsed = parse_resume_text(raw_text, target_role)
    return Response(parsed)


@api_view(['POST'])
def upload_parse_and_save(request):
    """
    Upload + parse + immediately save to DB.
    Returns the created ResumeProfile with full nested data.
    """
    file = request.FILES.get('file')
    if not file:
        return Response({'error': 'No file uploaded'}, status=400)

    target_role = request.data.get('target_role', '')
    raw_text    = ''

    if file.name.endswith('.pdf'):
        with pdfplumber.open(BytesIO(file.read())) as pdf:
            raw_text = '\n'.join(p.extract_text() or '' for p in pdf.pages)
    elif file.name.endswith('.docx'):
        from docx import Document
        doc      = Document(BytesIO(file.read()))
        raw_text = '\n'.join(p.text for p in doc.paragraphs)
    else:
        raw_text = file.read().decode('utf-8')

    parsed = parse_resume_text(raw_text, target_role)

    # Save everything
    resume = ResumeProfile.objects.create(
        full_name   = parsed.get('full_name', ''),
        email       = parsed.get('email', ''),
        phone       = parsed.get('phone', ''),
        summary     = parsed.get('summary', ''),
        target_role = parsed.get('target_role', target_role),
    )

    for exp in parsed.get('work_experiences', []):
        WorkExperience.objects.create(resume=resume, **exp)

    for edu in parsed.get('education', []):
        Education.objects.create(resume=resume, **edu)

    for proj in parsed.get('projects', []):
        Project.objects.create(resume=resume, **proj)

    for skill in parsed.get('skills', []):
        Skill.objects.create(resume=resume, **skill)

    for cert in parsed.get('certifications', []):
        Certification.objects.create(resume=resume, **cert)

    return Response(ResumeProfileSerializer(resume).data, status=201)


# ── VERSION HISTORY ────────────────────────────────────────────────────────

@api_view(['GET'])
def version_history(request, resume_pk):
    try:
        resume = ResumeProfile.objects.get(pk=resume_pk)
    except ResumeProfile.DoesNotExist:
        return Response({'error': 'Resume not found'}, status=404)

    versions = resume.versions.all()
    return Response(ResumeVersionSerializer(versions, many=True).data)


# ── EXPORT ─────────────────────────────────────────────────────────────────

@api_view(['GET'])
def export_json(request, resume_pk):
    try:
        resume = ResumeProfile.objects.get(pk=resume_pk)
    except ResumeProfile.DoesNotExist:
        return Response({'error': 'Resume not found'}, status=404)

    data = ResumeProfileSerializer(resume).data

    # Save version snapshot
    count = resume.versions.count() + 1
    ResumeVersion.objects.create(
        resume   = resume,
        label    = f"v{count}",
        snapshot = data,
        template = 'json'
    )

    response = HttpResponse(
        json.dumps(data, indent=2),
        content_type='application/json'
    )
    response['Content-Disposition'] = f'attachment; filename="{resume.full_name}_resume.json"'
    return response